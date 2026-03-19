#!/usr/bin/env python3
"""Build styled .docx documentation for the Secret Server Migration system.

Table style: dark navy headers (#1A365D) with white text, alternating row shading (#F7FAFC),
secondary tables with light blue headers (#E8F4F8). Arial font, 0.75in margins.

Matches the APM Terminals format used by the CyberArk migration documentation.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re

# ── Colors ───────────────────────────────────────────────────
NAVY = "1A365D"
DARK_BLUE = "2C5282"
MED_BLUE = "2B6CB0"
WHITE = "FFFFFF"
LIGHT_GRAY = "F7FAFC"
LIGHT_BLUE = "E8F4F8"
BODY_TEXT = "333333"
SUBTLE = "666666"
RED_ACCENT = "C53030"


def keep_with_next(paragraph):
    """Set 'Keep with next' on a paragraph so it stays on the same page as the following element."""
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:keepNext {nsdecls("w")}/>'))


def keep_lines_together(paragraph):
    """Set 'Keep lines together' so the paragraph doesn't split across pages."""
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:keepLines {nsdecls("w")}/>'))


def set_row_cant_split(row):
    """Prevent a table row from splitting across a page break."""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    margins = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(margins)


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def add_styled_table(doc, headers, rows, header_bg=NAVY, header_fg=WHITE, alt_shading=True):
    """Add a table with APM Terminals styling."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    # Header row
    header_row = table.rows[0]
    set_row_cant_split(header_row)
    trPr = header_row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

    for i, header_text in enumerate(headers):
        cell = header_row.cells[i]
        set_cell_shading(cell, header_bg)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        keep_lines_together(p)
        keep_with_next(p)
        run = p.add_run(header_text)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(10)
        if header_fg:
            run.font.color.rgb = RGBColor.from_string(header_fg)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        set_row_cant_split(row)
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_margins(cell)
            if alt_shading and r_idx % 2 == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            keep_lines_together(p)
            if r_idx == 0:
                keep_with_next(p)
            run = p.add_run(str(cell_text))
            run.font.name = "Arial"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(BODY_TEXT)

    return table


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    keep_with_next(h)
    keep_lines_together(h)
    if level == 1:
        # Each major section starts on a new page — prevents orphaned headings at bottom of page
        pPr = h._p.get_or_add_pPr()
        pPr.append(parse_xml(f'<w:pageBreakBefore {nsdecls("w")}/>'))
    for run in h.runs:
        run.font.name = "Arial"
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor.from_string(NAVY)
        elif level == 2:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
        elif level == 3:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor.from_string(MED_BLUE)
    return h


def add_para(doc, text, bold=False, size=10, color=BODY_TEXT, space_after=Pt(4), keep_next=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    p.paragraph_format.space_after = space_after
    if keep_next:
        keep_with_next(p)
        keep_lines_together(p)
    return p


def add_bullet(doc, text, bold=False, keep_next=True):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(BODY_TEXT)
    if keep_next:
        keep_with_next(p)
    return p


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    keep_lines_together(p)
    keep_with_next(p)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string("2D3748")
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EDF2F7" w:val="clear"/>')
    p._p.get_or_add_pPr().append(shading)
    return p


def build():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    # ── Title Block ──────────────────────────────────────────
    p = doc.add_paragraph()
    run = p.add_run("CyberArk PAM Migration")
    run.font.name = "Arial"
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY)

    p = doc.add_paragraph()
    run = p.add_run("8-Agent AI Orchestrator System")
    run.font.name = "Arial"
    run.font.size = Pt(20)
    run.bold = True

    p = doc.add_paragraph()
    run = p.add_run("CyberArk On-Premises \u2192 Delinea Secret Server")
    run.font.name = "Arial"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string(SUBTLE)

    p = doc.add_paragraph()
    run = p.add_run("Technical Documentation")
    run.font.name = "Arial"
    run.font.size = Pt(16)

    doc.add_paragraph()  # Spacer

    # ══════════════════════════════════════════════════════════
    # SECTION: Architecture
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "Architecture", 1)

    add_styled_table(doc,
        ["Component", "File", "Description"],
        [
            ["Orchestrator", "coordinator.py", "Sequences agents per phase, signal handlers, state management"],
            ["CLI", "cli.py", "User interface (status/start/run/advance/agent/preflight/agents)"],
            ["Agent Base", "core/base.py", "AgentBase ABC + AgentResult dataclass (status validation, approval timeout)"],
            ["State Machine", "core/state.py", "Atomic writes, fcntl file locking, backup recovery, raw data separation"],
            ["Audit Logger", "core/logging.py", "SHA-256 hash chain, SIEM-ready JSONL format"],
            ["On-Prem Client", "core/cyberark_client.py", "CyberArk PVWA REST client (context manager, env var credentials)"],
            ["SS Client", "core/secret_server_client.py", "Delinea Secret Server REST client (OAuth2, folder/secret/template ops)"],
            [".gitignore", ".gitignore", "Excludes config.json, credentials, output/, __pycache__/"],
        ]
    )

    # ══════════════════════════════════════════════════════════
    # SECTION: Architecture Mapping
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "CyberArk \u2192 Secret Server Architecture Mapping", 1)

    add_para(doc, "The target system (Delinea Secret Server) has a fundamentally different data model from CyberArk. This is NOT a simple URL swap \u2014 it is a full re-architecture.", keep_next=True)

    add_styled_table(doc,
        ["CyberArk Concept", "Secret Server Equivalent", "Migration Impact"],
        [
            ["Safe", "Folder", "Hierarchical (not flat). Permissions inherited from parent."],
            ["Platform", "Secret Template", "Must map or rebuild. Different field structures."],
            ["Account", "Secret", "Different JSON schema. Field slug mapping required."],
            ["CPM", "RPC / Distributed Engine", "Plugin rewrite required. RPC is per-template."],
            ["PSM", "Protocol Handler / SSH Proxy", "NO migration path for session recordings."],
            ["CCP/AAM", "REST API /api/v1/secrets/{id} + OAuth2", "Full re-architecture of all integrations."],
            ["22 safe permissions", "4 roles (Owner, Edit, View, List)", "LOSSY \u2014 permission collapse with escalation risk."],
            ["/PasswordVault/api/", "/api/v1/", "Completely different API surface."],
        ]
    )

    # ══════════════════════════════════════════════════════════
    # SECTION: Agent Registry
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "Agent Registry", 1)

    add_styled_table(doc,
        ["Agent", "Class", "Phases", "Key Capabilities"],
        [
            ["01-discovery", "DiscoveryAgent", "P1", "Applications API, multi-signal NHI detection, integration scanning"],
            ["02-gap-analysis", "GapAnalysisAgent", "P1", "10-domain assessment, SS template coverage, permission model risk"],
            ["03-permissions", "PermissionMappingAgent", "P1, P3", "22\u21924 lossy role translation, escalation detection, loss tracking"],
            ["04-etl", "ETLOrchestrationAgent", "P4, P5", "Safe\u2192Folder, Platform\u2192Template, Account\u2192Secret, watchdog timer"],
            ["05-heartbeat", "HeartbeatAgent", "P4-P6", "10 validation checks (heartbeat, permissions, RPC, audit continuity)"],
            ["06-integration", "IntegrationRepointingAgent", "P5, P6", "CCP/AAM code scanning, SS REST API replacement code generation"],
            ["07-compliance", "ComplianceAgent", "P5, P7", "PCI-DSS, NIST 800-53, HIPAA, SOX + SS-specific risk tracking"],
            ["08-runbook", "RunbookAgent", "All", "Phase gate management, human approval workflow"],
        ]
    )

    # ══════════════════════════════════════════════════════════
    # SECTION: Phase Execution Map
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "Phase Execution Map", 1)

    add_styled_table(doc,
        ["Phase", "Name", "Agents", "Human Gate"],
        [
            ["P0", "Environment Setup", "(manual)", "\u2014"],
            ["P1", "Discovery & Dependency Mapping", "01, 02, 03", "Review discovery + gaps + permission loss report (22\u21924)"],
            ["P2", "Infrastructure Preparation", "(manual)", "\u2014"],
            ["P3", "Folder & Template Migration", "03", "Approve folder structure and permission mapping plan"],
            ["P4", "Pilot Migration", "04, 05", "Approve pilot results before production waves"],
            ["P5", "Production Batches", "04, 05, 06, 07", "Approve all production batch results"],
            ["P6", "Parallel Running & Cutover", "05, 06, 07", "Approve cutover (decommission CyberArk read-only)"],
            ["P7", "Decommission & Close-Out", "07", "Final sign-off (confirm CyberArk audit archive complete)"],
        ]
    )

    # ══════════════════════════════════════════════════════════
    # SECTION: Credential Management
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "Credential Management", 1)

    add_para(doc, "Credentials should be provided via environment variables (recommended) rather than stored in config.json.", keep_next=True)

    add_styled_table(doc,
        ["Environment Variable", "Purpose"],
        [
            ["CYBERARK_USERNAME", "On-prem PVWA service account username"],
            ["CYBERARK_PASSWORD", "On-prem PVWA service account password"],
            ["SS_CLIENT_ID", "Secret Server OAuth2 client ID"],
            ["SS_CLIENT_SECRET", "Secret Server OAuth2 client secret"],
            ["SS_USERNAME", "Secret Server username (legacy auth only)"],
            ["SS_PASSWORD", "Secret Server password (legacy auth only)"],
        ]
    )

    add_para(doc, "Passwords are zeroed from memory immediately after the authentication attempt. On token expiry (HTTP 401), credentials are re-read from environment variables for re-authentication. Proactive re-auth triggers when the token is within 60 seconds of expiry.", size=9)

    # ══════════════════════════════════════════════════════════
    # SECTION: Authentication
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "Authentication", 1)

    add_heading(doc, "CyberArk On-Prem (Source)", 2)
    add_styled_table(doc,
        ["Auth Type", "Endpoint", "Notes"],
        [
            ["CyberArk", "/Auth/CyberArk/Logon", "Local vault authentication"],
            ["LDAP", "/Auth/LDAP/Logon", "Active Directory (recommended)"],
            ["RADIUS", "/Auth/RADIUS/Logon", "Multi-factor authentication"],
            ["Windows", "/Auth/Windows/Logon", "Windows integrated auth"],
            ["SAML", "\u2014", "Not supported (requires browser redirect flow)"],
        ]
    )

    add_heading(doc, "Secret Server (Target) \u2014 OAuth2 (Recommended)", 2)
    add_para(doc, "Modern Secret Server installations use OAuth2 client credentials:", keep_next=True)
    add_code_block(doc, "Token URL: {base_url}/oauth2/token\nGrant type: client_credentials\nFormat: https://secretserver.company.com/SecretServer")

    add_heading(doc, "Secret Server (Legacy)", 2)
    add_para(doc, "Set auth_method to 'legacy' in config. Uses password grant type at the same /oauth2/token endpoint with username/password instead of client credentials.")

    # ══════════════════════════════════════════════════════════
    # SECTION: Configuration
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "Configuration", 1)

    add_heading(doc, "config.json", 2)
    add_styled_table(doc,
        ["Key", "Purpose", "Default"],
        [
            ["cyberark_on_prem.base_url", "On-prem PVWA URL", "\u2014"],
            ["cyberark_on_prem.auth_type", "CyberArk / LDAP / RADIUS / Windows", "LDAP"],
            ["cyberark_on_prem.verify_ssl", "SSL certificate verification", "true"],
            ["cyberark_on_prem.timeout", "Request timeout (seconds)", "30"],
            ["cyberark_on_prem.rate_limit", "Seconds between requests", "0.1"],
            ["cyberark_on_prem.batch_size", "Pagination size (max 1000)", "1000"],
            ["secret_server.base_url", "Secret Server URL", "\u2014"],
            ["secret_server.auth_method", "oauth2 (modern) or legacy", "oauth2"],
            ["secret_server.batch_size", "Import pagination size (max 1000)", "500"],
            ["secret_server.default_folder_id", "Root folder ID for imports", "-1"],
            ["secret_server.inherit_permissions", "Folders inherit parent permissions", "true"],
            ["servicenow.instance_url", "ServiceNow instance URL", "\u2014"],
            ["output_dir", "Output directory", "./output"],
            ["log_level", "Logging level", "INFO"],
            ["environment", "Environment context", "dev"],
        ]
    )

    # ══════════════════════════════════════════════════════════
    # SECTION: Security Features
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "Security Features", 1)

    add_styled_table(doc,
        ["Feature", "Implementation"],
        [
            ["Secrets Management", "Env var credentials, passwords zeroed from memory after auth, error sanitization strips secrets"],
            ["Atomic State Writes", "temp file \u2192 fsync \u2192 os.replace (POSIX atomic) + fcntl.flock(LOCK_EX) + .bak backup"],
            ["Signal Handlers", "SIGTERM/SIGINT save state and exit cleanly (exit code 128 + signal)"],
            ["Tamper-Evident Logging", "SHA-256 hash chain in every audit event \u2014 tampering breaks the chain"],
            ["Watchdog Timer", "threading.Timer auto-unfreezes all frozen accounts on ETL timeout (default 120 min)"],
            ["Approval Timeouts", "30-minute default, auto-deny in non-interactive mode (no TTY)"],
            ["Proactive Token Refresh", "Re-authenticates when token is within 60 seconds of expiry"],
            ["Error Sanitization", "_safe_error() strips passwords, secrets, tokens from HTTP error messages"],
            [".gitignore", "Excludes config.json, *.pem, *.key, .env, output/, __pycache__/"],
        ],
        header_bg=LIGHT_BLUE,
        header_fg=None,
        alt_shading=False,
    )

    # ══════════════════════════════════════════════════════════
    # SECTION: Permission Translation (Agent 03) — THE BIG SECTION
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "Permission Translation (Agent 03) \u2014 LOSSY", 1)

    add_para(doc,
        "CyberArk uses 22 individual safe member permissions. Secret Server uses a 4-tier role model "
        "(Owner, Edit, View, List). This translation is inherently LOSSY \u2014 9 CyberArk permissions "
        "have NO Secret Server equivalent, and some members will receive MORE access than they had before.",
        bold=True, color=RED_ACCENT, keep_next=True,
    )

    add_heading(doc, "All 22 CyberArk Safe Permissions", 2)
    add_code_block(doc,
        "UseAccounts              RetrieveAccounts         ListAccounts\n"
        "AddAccounts              UpdateAccountContent     UpdateAccountProperties\n"
        "InitiateCPMAccountManagementOperations            SpecifyNextAccountContent\n"
        "RenameAccounts           DeleteAccounts           UnlockAccounts\n"
        "ManageSafe               ManageSafeMembers        BackupSafe\n"
        "ViewAuditLog             ViewSafeMembers          AccessWithoutConfirmation\n"
        "CreateFolders            DeleteFolders            MoveAccountsAndFolders\n"
        "RequestsAuthorizationLevel1                       RequestsAuthorizationLevel2"
    )

    add_heading(doc, "Secret Server 4-Tier Role Model", 2)
    add_styled_table(doc,
        ["SS Role", "Grants", "CyberArk Trigger Permissions"],
        [
            ["Owner", "Full admin \u2014 manage folder and member access",
             "Requires BOTH ManageSafe AND ManageSafeMembers"],
            ["Edit", "Create/modify/delete secrets in the folder",
             "Any of: AddAccounts, UpdateAccountContent, UpdateAccountProperties, "
             "DeleteAccounts, RenameAccounts, UnlockAccounts"],
            ["View", "View and retrieve secret values",
             "Any of: UseAccounts, RetrieveAccounts"],
            ["List", "See that secrets exist (no value access)",
             "Any of: ListAccounts, ViewSafeMembers, ViewAuditLog"],
        ]
    )

    add_heading(doc, "Permissions with NO Secret Server Equivalent (Always Lost)", 2)
    add_styled_table(doc,
        ["CyberArk Permission", "Why It's Lost"],
        [
            ["AccessWithoutConfirmation", "SS uses Workflow templates instead of dual-control bypass"],
            ["SpecifyNextAccountContent", "SS RPC sets passwords automatically \u2014 no manual setter"],
            ["BackupSafe", "Folder backup is a server-level operation in SS"],
            ["CreateFolders", "SS folder permissions are inherited, not per-member"],
            ["DeleteFolders", "SS folder permissions are inherited, not per-member"],
            ["MoveAccountsAndFolders", "SS move is a separate admin permission, not folder-level"],
            ["RequestsAuthorizationLevel1", "SS uses Workflow templates \u2014 different approval model"],
            ["RequestsAuthorizationLevel2", "SS uses Workflow templates \u2014 different approval model"],
            ["InitiateCPMAccountManagementOperations", "SS RPC is all-or-nothing per template, not per-member"],
        ]
    )

    add_heading(doc, "Escalation Risks (Over-Provisioning)", 2)
    add_para(doc, "The 22\u21924 collapse can grant members MORE access than they originally had:", keep_next=True)
    add_styled_table(doc,
        ["Scenario", "CyberArk Original", "SS Result", "Risk"],
        [
            ["View\u2192Edit escalation",
             "UseAccounts + RetrieveAccounts + UnlockAccounts",
             "Edit role (full create/modify/delete)",
             "Member gets secret creation and deletion \u2014 only had view + unlock"],
            ["Admin-only\u2192Full Owner",
             "ManageSafe + ManageSafeMembers (no data permissions)",
             "Owner role (full secret access)",
             "Admin who could not see passwords now has full access to all secrets"],
        ],
        header_bg=LIGHT_BLUE,
        header_fg=None,
    )

    add_para(doc,
        "Agent 03 produces a detailed loss report per member showing: assigned SS role, "
        "lost permissions, and escalation risk flags. This report MUST be reviewed by a "
        "human before proceeding to Phase P3.",
        size=9,
    )

    # ══════════════════════════════════════════════════════════
    # SECTION: ETL Pipeline (Agent 04)
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "ETL Pipeline (Agent 04)", 1)

    add_para(doc, "Each batch executes 7 steps in sequence with real API calls:", keep_next=True)

    add_styled_table(doc,
        ["Step", "Action", "API Call"],
        [
            ["1. FREEZE", "Disable CPM automatic management",
             "PATCH /Accounts/{id} (automaticManagementEnabled=false)"],
            ["2. EXPORT", "Pull full account details + retrieve passwords",
             "GET /Accounts/{id} + POST .../Password/Retrieve"],
            ["3. TRANSFORM", "Map CyberArk fields to Secret Server schema",
             "Local (userName\u2192username, address\u2192machine, platformId\u2192template)"],
            ["4. FOLDER CREATION", "Create folder hierarchy in Secret Server",
             "POST /api/v1/folders (parent: /Imported/{SafeName})"],
            ["5. IMPORT", "Push secrets with retry + rate limiting",
             "POST /api/v1/secrets (rate-limited)"],
            ["6. HEARTBEAT", "Trigger password verification via RPC",
             "POST /api/v1/secrets/{id}/heartbeat"],
            ["7. UNFREEZE", "Re-enable CPM management on CyberArk",
             "PATCH /Accounts/{id} (automaticManagementEnabled=true)"],
        ]
    )

    add_heading(doc, "Field Mapping (CyberArk \u2192 Secret Server)", 2)
    add_styled_table(doc,
        ["CyberArk Field", "Secret Server Field", "Notes"],
        [
            ["name", "name", "Secret display name"],
            ["userName", "username (slug)", "Template field item"],
            ["address", "machine (slug)", "Template field item"],
            ["_password", "password (slug)", "Retrieved via CyberArk API, set via template field"],
            ["safeName", "folderId", "Mapped via folder creation step"],
            ["platformId", "secretTemplateId", "Mapped via platform\u2192template lookup"],
            ["platformAccountProperties", "notes (slug)", "Preserved as key: value text in notes field"],
        ],
        header_bg=LIGHT_BLUE,
        header_fg=None,
    )

    add_heading(doc, "Implementation Details", 2)
    add_bullet(doc, "Password retrieval: Accounts where retrieval fails are skipped (not imported with empty secrets)")
    add_bullet(doc, "Folder hierarchy: Parent 'Imported' folder created at root, then per-safe child folders")
    add_bullet(doc, "Conflict handling: HTTP 409 on import treated as idempotent success (already exists)")
    add_bullet(doc, "Failure threshold: >10% batch failure = 'failed', otherwise 'partial'")
    add_bullet(doc, "Both clients use context managers (with ... as client:) to prevent connection leaks")
    add_bullet(doc, "Emergency unfreeze: On any exception, all frozen accounts are automatically unfrozen", keep_next=False)

    add_heading(doc, "Wave Classification (5 Waves by Risk)", 2)
    add_styled_table(doc,
        ["Wave", "Risk", "Accounts"],
        [
            ["1", "LOW", "Test/Dev/Sandbox"],
            ["2", "MEDIUM", "Standard user accounts"],
            ["3", "MEDIUM-HIGH", "Infrastructure/network/admin"],
            ["4", "HIGH", "NHIs without CCP/AAM"],
            ["5", "CRITICAL", "NHIs with CCP/AAM (uses Applications API data)"],
        ]
    )

    # ══════════════════════════════════════════════════════════
    # SECTION: Platform → Template Mapping
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "Platform \u2192 Template Mapping", 1)

    add_para(doc, "CyberArk Platforms must be mapped to Secret Server Templates before migration. The ETL agent uses this mapping during the transform step.", keep_next=True)

    add_styled_table(doc,
        ["CyberArk Platform", "Secret Server Template", "Notes"],
        [
            ["WinServerLocal", "Windows Account", "Local administrator accounts"],
            ["WinDomain", "Active Directory Account", "Domain service accounts"],
            ["WinServiceAccount", "Windows Service Account", "Windows service credentials"],
            ["UnixSSH", "Unix Account (SSH)", "Standard SSH password auth"],
            ["UnixSSHKeys", "Unix Account (SSH Key Rotation)", "SSH key-based authentication"],
            ["Oracle", "Oracle Account", "Oracle database credentials"],
            ["MSSql", "SQL Server Account", "Microsoft SQL Server"],
            ["MySQL", "MySQL Account", "MySQL database credentials"],
            ["AzureServicePrincipal", "Azure Service Principal", "Azure AD app registrations"],
            ["AWSAccessKeys", "Amazon IAM Key", "AWS IAM access keys"],
            ["(custom)", "(custom template required)", "Must create SS template before migration"],
        ]
    )

    add_para(doc, "Custom CyberArk platforms that do not appear in this mapping require manual Secret Server template creation. Agent 02 (Gap Analysis) identifies these gaps during Phase P1.", size=9)

    # ══════════════════════════════════════════════════════════
    # SECTION: Validation Checks (Agent 05)
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "Validation Checks (Agent 05)", 1)

    add_styled_table(doc,
        ["#", "Check", "Description", "SS-Specific Notes"],
        [
            ["1", "Count comparison", "Source vs target counts within threshold", "CyberArk accounts vs SS secrets"],
            ["2", "Heartbeat status", "All imported secrets have valid credentials", "POST /api/v1/secrets/{id}/heartbeat"],
            ["3", "Permission mapping", "Translations applied, exceptions flagged", "Flags 22\u21924 escalation risks"],
            ["4", "Folder structure", "Safe\u2192Folder hierarchy preserved", "Verifies /Imported/{SafeName} tree"],
            ["5", "Metadata integrity", "Descriptions and custom fields intact", "Properties preserved in notes field"],
            ["6", "Group assignments", "User/group folder permission mappings", "Verifies folder permission roles"],
            ["7", "Password policies", "Rotation policies applied in target", "WARNING: RPC requires manual config per template"],
            ["8", "Access patterns", "No unexpected permission escalations", "Flags >30% Edit/Owner as over-provisioned"],
            ["9", "Audit continuity", "Audit trail preserved across migration", "WARNING: CyberArk logs do NOT transfer to SS"],
            ["10", "Recording preservation", "PSM recordings archived", "WARNING: PSM recordings cannot migrate to SS"],
        ]
    )

    # ══════════════════════════════════════════════════════════
    # SECTION: Integration Repointing (Agent 06)
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "Integration Repointing (Agent 06)", 1)

    add_para(doc, "CyberArk CCP/AAM integrations require a FULL re-architecture to use Secret Server's REST API. This is not a URL swap \u2014 the authentication model (OAuth2 vs AppID), endpoint structure, and data format are completely different.", keep_next=True)

    add_styled_table(doc,
        ["CyberArk Pattern", "Secret Server Replacement"],
        [
            ["GET /AIMWebService/api/Accounts?AppID=X&Safe=Y&Object=Z",
             "POST /oauth2/token + GET /api/v1/secrets/{id}/fields/password"],
            ["AppID-based authentication", "OAuth2 client credentials (Bearer token)"],
            ["psPAS / Get-PASAccount (PowerShell)", "Thycotic.SecretServer module / Get-TssSecret"],
            ["CyberArk.AIM.NetPasswordSDK (.NET)", "HttpClient + OAuth2 + /api/v1/secrets"],
            ["com.cyberark.aim (Java)", "HttpClient + OAuth2 + /api/v1/secrets"],
        ],
        header_bg=LIGHT_BLUE,
        header_fg=None,
    )

    add_para(doc, "Agent 06 scans configured directories for CyberArk CCP patterns and generates language-specific replacement code templates (Python, PowerShell, C#, Java).", size=9)

    # ══════════════════════════════════════════════════════════
    # SECTION: Compliance Frameworks (Agent 07)
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "Compliance Frameworks (Agent 07)", 1)

    add_para(doc, "Maps migration actions to controls in 4 frameworks using data-driven check definitions (no lambdas \u2014 serializable config):", keep_next=True)

    add_styled_table(doc,
        ["Framework", "Control Groups"],
        [
            ["PCI-DSS v4.0", "Access control (8.x), Audit trail (10.x), Change management (6.x), Permission integrity (22\u21924 loss risk)"],
            ["NIST 800-53 Rev5", "IA-2/4/5, AC-2/3/6, AU-2/3/6/11, CA-7, SI-4"],
            ["HIPAA Security Rule", "164.312(a)(1), 164.312(b), 164.312(c)(1), 164.312(d)"],
            ["SOX IT Controls", "CC6.1-3, CC7.1-2, CC8.1"],
        ]
    )

    # ══════════════════════════════════════════════════════════
    # SECTION: SS-Specific Compliance Risks
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "Secret Server \u2014 Specific Compliance Risks", 1)

    add_para(doc, "Three compliance risks are unique to the Secret Server migration path and must be documented for auditors:", keep_next=True)

    add_styled_table(doc,
        ["Risk", "Severity", "Description", "Mitigation"],
        [
            ["Audit Log Discontinuity", "HIGH",
             "CyberArk audit history does NOT transfer to Secret Server. "
             "There will be a gap in the audit trail during migration.",
             "Maintain CyberArk in read-only mode for audit retention period. "
             "Document the gap for auditors."],
            ["Permission Model Simplification", "HIGH",
             "CyberArk's 22 granular permissions collapse to 4 SS roles. "
             "Some users may receive more access than they had before (escalation risk).",
             "Review Agent 03 permission loss report. Document all escalations. "
             "Consider SS Workflow templates for dual-control equivalence."],
            ["PSM Session Recording Loss", "MEDIUM",
             "PSM session recordings cannot be migrated to Secret Server.",
             "Archive all recordings before CyberArk decommission. "
             "Maintain read-only CyberArk access for recording review."],
        ]
    )

    # ══════════════════════════════════════════════════════════
    # SECTION: NHI Detection
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "NHI Detection (Multi-Signal)", 1)

    add_styled_table(doc,
        ["Signal", "Priority", "Method"],
        [
            ["Platform-based", "Highest", "UnixSSHKeys, WinServiceAccount, WinScheduledTask, AzureServicePrincipal, AWSAccessKeys, HashiCorpVault"],
            ["Name patterns", "Medium", "^svc[_-], ^app[_-], ^api[_-], ^bot[_-], ^sys[_-], service.?account, daemon, scheduler"],
            ["Safe name patterns", "Lower", "appcred, servicecred, automation, cicd, pipeline, appidentit, machineidentit"],
        ],
        header_bg=LIGHT_BLUE,
        header_fg=None,
    )

    # ══════════════════════════════════════════════════════════
    # SECTION: API Endpoints
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "API Endpoints", 1)

    add_heading(doc, "CyberArk On-Prem PVWA (Source)", 2)
    add_styled_table(doc,
        ["Endpoint", "Method", "Purpose"],
        [
            ["/Auth/{type}/Logon", "POST", "Authentication (CyberArk/LDAP/RADIUS/Windows)"],
            ["/Auth/Logoff", "POST", "Session termination"],
            ["/Accounts", "GET", "Account enumeration (paginated, count-based)"],
            ["/Accounts/{id}", "GET", "Full account details (includes linked accounts)"],
            ["/Accounts/{id}", "PATCH", "Update account properties (JSON Patch)"],
            ["/Accounts/{id}/Password/Retrieve", "POST", "Password retrieval"],
            ["/Safes", "GET", "Safe enumeration (system safes filtered)"],
            ["/Safes/{name}/Members", "GET", "Safe member/permission enumeration"],
            ["/Platforms/Targets", "GET", "Platform enumeration (v12+)"],
            ["/Applications", "GET", "CCP/AAM application enumeration"],
            ["/Applications/{id}/Authentications", "GET", "Application auth methods"],
            ["/Activities", "GET", "Audit log retrieval"],
            ["/Server/Verify", "GET", "Version and health info"],
        ]
    )

    add_heading(doc, "Delinea Secret Server (Target)", 2)
    add_styled_table(doc,
        ["Endpoint", "Method", "Purpose"],
        [
            ["/oauth2/token", "POST", "OAuth2 token (client_credentials or password grant)"],
            ["/api/v1/folders", "GET", "Folder enumeration (hierarchical)"],
            ["/api/v1/folders", "POST", "Folder creation (with parent ID + permission inheritance)"],
            ["/api/v1/folder-permissions", "POST", "Set folder permission (folderAccessRoleName + secretAccessRoleName)"],
            ["/api/v1/folder-permissions/{id}", "PUT", "Update folder permission"],
            ["/api/v1/secrets", "GET", "Secret enumeration (paginated, take/skip)"],
            ["/api/v1/secrets", "POST", "Secret creation (template ID + field items)"],
            ["/api/v1/secrets/{id}", "GET", "Full secret details"],
            ["/api/v1/secrets/{id}/fields/password", "GET", "Password retrieval"],
            ["/api/v1/secrets/{id}/fields/password", "PUT", "Password update"],
            ["/api/v1/secrets/{id}/heartbeat", "POST", "Trigger heartbeat (password verification via RPC)"],
            ["/api/v1/secrets/{id}/state", "GET", "Heartbeat/RPC status"],
            ["/api/v1/secrets/{id}", "DELETE", "Secret deletion (rollback)"],
            ["/api/v1/secret-templates", "GET", "Template enumeration"],
            ["/api/v1/secret-templates", "POST", "Template creation"],
            ["/api/v1/users", "GET", "User enumeration"],
            ["/api/v1/groups", "GET", "Group enumeration"],
            ["/api/v1/distributed-engine/sites", "GET", "Distributed engine site enumeration"],
            ["/api/v1/version", "GET", "Server version and health"],
        ]
    )

    # ══════════════════════════════════════════════════════════
    # SECTION: State Machine
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "State Machine", 1)

    add_para(doc, "Migration state persists to output/state/migration_state.json via atomic writes:", keep_next=True)

    add_styled_table(doc,
        ["Feature", "Implementation"],
        [
            ["Persistence", "JSON file with atomic writes (temp \u2192 fsync \u2192 os.replace)"],
            ["Concurrency", "fcntl.flock(LOCK_EX) prevents concurrent writes"],
            ["Backup", ".bak file created before every write; recovery on corruption"],
            ["Raw Data", "Large datasets stored in output/state/raw/ (not in main state file)"],
            ["List Caps", "Steps: 5000, Errors: 1000, Approvals: 500"],
            ["Resume", "Coordinator picks up from last completed step after restart"],
        ],
        header_bg=LIGHT_BLUE,
        header_fg=None,
        alt_shading=False,
    )

    # ══════════════════════════════════════════════════════════
    # SECTION: Test Environment
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "Test Environment & Sandbox (Next Steps Decision Point)", 1)

    add_para(doc, "The system currently has no dedicated virtual test environment. Before running the ETL pipeline for the first time, a decision is needed on the testing approach.", keep_next=True)

    add_heading(doc, "Current Capabilities", 2)
    add_styled_table(doc,
        ["Capability", "What It Does", "Limitations"],
        [
            ["--dry-run mode", "Skips agent execution, prints what would run", "No data flow simulation"],
            ["environment config", "Stamps dev/staging/prod into audit logs", "Metadata only \u2014 no behavior change"],
            ["P4 Pilot (Wave 1)", "Limits first batch to test/dev accounts only", "Hits real APIs \u2014 not a sandbox"],
        ],
        header_bg=LIGHT_BLUE,
        header_fg=None,
    )

    add_heading(doc, "Decision Matrix", 2)
    add_styled_table(doc,
        ["Option", "Description", "Effort", "Fidelity"],
        [
            ["A. Lab Instance", "CyberArk PVWA dev/lab + Secret Server test instance", "Low", "High"],
            ["B. Mock API Server", "Flask/FastAPI emulating PVWA + SS APIs", "Medium", "Medium"],
            ["C. Unit Test Fixtures", "pytest with mocked API clients", "Medium", "Medium"],
            ["D. Docker Harness", "Mock API containers with pre-seeded data", "High", "High"],
            ["E. Hybrid (Recommended)", "A for integration + C for CI/CD regression", "Medium", "High"],
        ]
    )

    add_heading(doc, "Recommended: Hybrid (Option E)", 2)
    add_bullet(doc, "Integration testing: Separate config files per environment (config.dev.json, config.prod.json)")
    add_bullet(doc, "CI/CD regression: pytest fixtures with mocked API responses")
    add_bullet(doc, "Wave 1 is the built-in canary \u2014 test/dev/sandbox accounts migrated first as a safe first-run", keep_next=False)

    # ══════════════════════════════════════════════════════════
    # SECTION: CLI Commands
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "CLI Commands", 1)

    add_styled_table(doc,
        ["Command", "Description"],
        [
            ["python3 cli.py status", "Show migration status and phase breakdown"],
            ["python3 cli.py start <id>", "Start new migration with given ID"],
            ["python3 cli.py run <P#>", "Run all agents for a phase"],
            ["python3 cli.py run <P#> --dry-run", "Simulate phase without API calls"],
            ["python3 cli.py advance", "Advance to next phase"],
            ["python3 cli.py agent <key>", "Run a single agent (e.g., 01-discovery)"],
            ["python3 cli.py agent <key> --phase P#", "Run agent for a specific phase"],
            ["python3 cli.py preflight", "Run all agent preflight checks"],
            ["python3 cli.py agents", "List all available agents"],
        ]
    )

    # ══════════════════════════════════════════════════════════
    # SECTION: Audit Logging
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "Audit Logging", 1)

    add_para(doc, "All agent actions are logged as structured JSON events (JSONL) to output/logs/<agent_id>.audit.jsonl. Every event includes a chain_hash \u2014 a running SHA-256 hash linking each event to its predecessor. Tampering with log entries breaks the chain, making modifications detectable during compliance audits.", keep_next=True)

    add_code_block(doc,
        '{\n'
        '  "timestamp": "2026-03-03T03:16:47.853Z",\n'
        '  "agent_id": "agent_01_discovery",\n'
        '  "environment": "prod",\n'
        '  "action": "discovery_start",\n'
        '  "details": {"phase": "P1"},\n'
        '  "user": "maverick",\n'
        '  "session_id": "6a410ecc",\n'
        '  "result": "success",\n'
        '  "chain_hash": "a3f8c2...e91d04"\n'
        '}'
    )

    # ══════════════════════════════════════════════════════════
    # SECTION: Migration Risk Summary
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "Migration Risk Summary", 1)

    add_para(doc, "Key risks specific to the CyberArk \u2192 Secret Server migration path, compared to Option A (Privilege Cloud):", keep_next=True)

    add_styled_table(doc,
        ["Risk Area", "Severity", "CyberArk \u2192 Privilege Cloud", "CyberArk \u2192 Secret Server"],
        [
            ["Permission model", "HIGH", "22 perms \u2192 22 perms (1:1)", "22 perms \u2192 4 roles (LOSSY)"],
            ["Audit continuity", "HIGH", "Logs can be migrated", "Logs do NOT transfer"],
            ["Session recordings", "MEDIUM", "PSM \u2192 Privilege Cloud recordings", "PSM recordings cannot migrate"],
            ["API surface", "MEDIUM", "Same REST API (/PasswordVault/api/)", "Completely different API (/api/v1/)"],
            ["CPM/RPC", "MEDIUM", "CPM plugins carry over", "RPC plugins must be rebuilt per template"],
            ["Integration rework", "HIGH", "CCP/AAM patterns similar", "Full re-architecture required (OAuth2)"],
        ]
    )

    # ── Save ─────────────────────────────────────────────────
    out_path = "/home/maverick/projects/iOPEX/Secret Server migration/Secret_Server_Migration_Agent_System_Documentation.docx"
    doc.save(out_path)
    print(f"Saved to {out_path}")

    # Also copy to OneDrive
    import shutil
    onedrive = "/mnt/c/Users/1/OneDrive/Documents/iOPEX/Secret_Server_Migration_Agent_System_Documentation.docx"
    try:
        shutil.copy2(out_path, onedrive)
        print(f"Copied to {onedrive}")
    except Exception as e:
        print(f"OneDrive copy failed (file may be open): {e}")
        print(f"Manual copy: cp '{out_path}' '{onedrive}'")


if __name__ == "__main__":
    build()
