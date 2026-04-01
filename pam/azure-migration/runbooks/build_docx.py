#!/usr/bin/env python3
"""
Build all migration runbooks as formatted .docx files.

Usage:
    python3 build_docx.py              # Build all runbooks
    python3 build_docx.py P1           # Build a specific runbook

Output: runbooks/docx/*.docx
"""

import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── iOPEX Brand Colors ─────────────────────────────────────────────────────
IOPEX_BLUE = RGBColor(0x00, 0x5B, 0xB5)      # Primary blue
IOPEX_DARK = RGBColor(0x1A, 0x1A, 0x2E)      # Dark navy
IOPEX_GRAY = RGBColor(0x6C, 0x75, 0x7D)      # Body gray
IOPEX_LIGHT = RGBColor(0xF4, 0xF6, 0xF8)     # Light bg
IOPEX_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
IOPEX_RED = RGBColor(0xDC, 0x35, 0x45)       # Warning/risk red
IOPEX_GREEN = RGBColor(0x28, 0xA7, 0x45)     # Success green
IOPEX_AMBER = RGBColor(0xFF, 0xC1, 0x07)     # Warning amber


# ── Helper: set cell background color ─────────────────────────────────────
def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    hex_color = '{:02X}{:02X}{:02X}'.format(rgb[0], rgb[1], rgb[2])
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_borders(table):
    """Apply subtle borders to all table cells."""
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'D0D7DE')
                tcBorders.append(border)
            tcPr.append(tcBorders)


def set_row_cant_split(row):
    """Prevent a table row from splitting across a page break."""
    tr = row._tr
    trPr = tr.find(qn('w:trPr'))
    if trPr is None:
        trPr = OxmlElement('w:trPr')
        tr.insert(0, trPr)
    for existing in trPr.findall(qn('w:cantSplit')):
        trPr.remove(existing)
    trPr.append(OxmlElement('w:cantSplit'))  # presence alone = cannot split


def set_row_header(row):
    """Mark row as repeating header when table spans pages."""
    tr = row._tr
    trPr = tr.find(qn('w:trPr'))
    if trPr is None:
        trPr = OxmlElement('w:trPr')
        tr.insert(0, trPr)
    if trPr.find(qn('w:tblHeader')) is None:
        trPr.append(OxmlElement('w:tblHeader'))


def set_keep_with_next(p):
    """Keep this paragraph on the same page as the next (for headings)."""
    pPr = p._p.get_or_add_pPr()
    keepNext = OxmlElement('w:keepNext')
    pPr.append(keepNext)


def set_keep_lines(p):
    """Keep all lines of this paragraph together (no mid-paragraph break)."""
    pPr = p._p.get_or_add_pPr()
    keepLines = OxmlElement('w:keepLines')
    pPr.append(keepLines)


def set_table_column_widths(table, col_count: int):
    """Set proportional column widths across the 6.5-inch content area."""
    content_width = Inches(6.5)
    col_width = content_width // col_count
    for row in table.rows:
        for cell in row.cells:
            cell.width = col_width


# ── Document Setup ─────────────────────────────────────────────────────────
def create_document() -> Document:
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.font.color.rgb = IOPEX_DARK

    return doc


# ── Cover / Title Block ────────────────────────────────────────────────────
def add_title_block(doc: Document, title: str, subtitle: str, meta: dict):
    # Blue header bar
    header_para = doc.add_paragraph()
    header_para.paragraph_format.space_before = Pt(0)
    header_para.paragraph_format.space_after = Pt(0)
    run = header_para.add_run("  iOPEX  |  SHIFT Migration System  |  CyberArk → KeeperPAM")
    run.font.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = IOPEX_WHITE
    header_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Shade the paragraph (approximation — header bar)
    pPr = header_para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '005BB5')
    pPr.append(shd)

    doc.add_paragraph()

    # Main title
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = IOPEX_BLUE
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Subtitle
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(subtitle)
    run.font.size = Pt(12)
    run.font.color.rgb = IOPEX_GRAY
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Meta table (scope, duration, etc.)
    if meta:
        table = doc.add_table(rows=len(meta), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        for i, (key, val) in enumerate(meta.items()):
            row = table.rows[i]
            set_row_cant_split(row)
            kc = row.cells[0]
            vc = row.cells[1]
            kc.width = Inches(1.4)
            vc.width = Inches(5.1)
            kc.paragraphs[0].add_run(key).font.bold = True
            kc.paragraphs[0].runs[0].font.color.rgb = IOPEX_BLUE
            kc.paragraphs[0].runs[0].font.size = Pt(9)
            vc.paragraphs[0].add_run(val).font.size = Pt(9)
            set_cell_bg(kc, IOPEX_LIGHT)
        set_cell_borders(table)

    # Divider
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("─" * 80)
    run.font.color.rgb = IOPEX_BLUE
    run.font.size = Pt(8)


# ── Heading Styles ─────────────────────────────────────────────────────────
def add_h1(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.page_break_before = False
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = IOPEX_BLUE

    # Blue underline via border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '005BB5')
    pBdr.append(bottom)
    pPr.append(pBdr)

    set_keep_with_next(p)


def add_h2(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = IOPEX_DARK
    set_keep_with_next(p)


def add_h3(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = IOPEX_BLUE
    set_keep_with_next(p)


# ── Body Text ─────────────────────────────────────────────────────────────
def add_body(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    text = text.strip()
    # Handle inline bold (**text**)
    parts = re.split(r'\*\*(.+?)\*\*', text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        run.font.size = Pt(10)
        run.font.color.rgb = IOPEX_DARK
        if i % 2 == 1:
            run.font.bold = True
    return p


def add_bullet(doc: Document, text: str, level: int = 0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    text = text.strip().lstrip('-').lstrip('*').strip()
    # Handle inline bold
    parts = re.split(r'\*\*(.+?)\*\*', text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        run.font.size = Pt(10)
        if i % 2 == 1:
            run.font.bold = True


def add_code_block(doc: Document, text: str):
    """Monospace code block with light gray background, kept together across pages."""
    lines = text.split('\n')
    paragraphs = []
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.right_indent = Inches(0.2)

        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F4F6F8')
        pPr.append(shd)

        run = p.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(8)
        run.font.color.rgb = IOPEX_DARK

        set_keep_lines(p)
        paragraphs.append(p)

    # Chain all lines except the last with keepNext so the block stays together
    for p in paragraphs[:-1]:
        set_keep_with_next(p)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ── Table Builder ──────────────────────────────────────────────────────────
def add_markdown_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return

    # Keep the paragraph before the table attached to the table —
    # prevents the preceding heading sitting alone at the bottom of a page
    if doc.paragraphs:
        set_keep_with_next(doc.paragraphs[-1])

    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Light List'

    col_width = Inches(6.5) // col_count
    for row in table.rows:
        for cell in row.cells:
            cell.width = col_width

    total_rows = len(rows)
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        set_row_cant_split(row)
        # Chain rows together so the whole table shifts as a unit if it fits on one page
        if i < total_rows - 1:
            for cell in row.cells:
                for para in cell.paragraphs:
                    set_keep_with_next(para)
        for j, cell_text in enumerate(row_data):
            if j >= col_count:
                break
            cell = row.cells[j]
            cell_text = cell_text.strip()
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)

            if i == 0:
                # Header row
                set_cell_bg(cell, IOPEX_BLUE)
                run = p.add_run(cell_text)
                run.font.bold = True
                run.font.color.rgb = IOPEX_WHITE
                run.font.size = Pt(9)
            else:
                # Data row — alternating shading
                if i % 2 == 0:
                    set_cell_bg(cell, IOPEX_LIGHT)
                # Handle inline bold **text** and inline code `text`
                parts = re.split(r'(\*\*.+?\*\*|`.+?`)', cell_text)
                for part in parts:
                    if not part:
                        continue
                    run = p.add_run(part.strip('*').strip('`') if part.startswith('**') or part.startswith('`') else part)
                    run.font.size = Pt(9)
                    if part.startswith('**') and part.endswith('**'):
                        run.font.bold = True
                    elif part.startswith('`') and part.endswith('`'):
                        run.font.name = 'Courier New'
                        run.font.color.rgb = IOPEX_BLUE

        row.height = None

    set_cell_borders(table)
    # Space after table
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)


# ── Checklist ─────────────────────────────────────────────────────────────
def add_checklist_item(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.2)
    run = p.add_run("☐  ")
    run.font.size = Pt(11)
    run.font.color.rgb = IOPEX_BLUE
    text_clean = re.sub(r'^[-\s]*\[.\]\s*', '', text.strip()).strip()
    run2 = p.add_run(text_clean)
    run2.font.size = Pt(10)
    run2.font.color.rgb = IOPEX_DARK


# ── Markdown Parser → docx ─────────────────────────────────────────────────
def parse_markdown_to_docx(doc: Document, md_path: str, title_meta: dict = None):
    content = Path(md_path).read_text(encoding='utf-8')
    lines = content.split('\n')

    i = 0
    in_code = False
    code_buffer = []
    in_table = False
    table_buffer = []

    # Extract title and subtitle from first 2 heading lines
    h1_text = ""
    h2_text = ""
    meta = title_meta or {}

    for line in lines:
        if line.startswith('# ') and not h1_text:
            h1_text = line[2:].strip()
        elif line.startswith('## ') and not h2_text:
            h2_text = line[3:].strip()
        if h1_text and h2_text:
            break

    add_title_block(doc, h1_text, h2_text, meta)

    skip_first_h1 = True
    skip_first_h2 = True

    while i < len(lines):
        line = lines[i]

        # Code block start/end
        if line.strip().startswith('```'):
            if in_code:
                add_code_block(doc, '\n'.join(code_buffer))
                code_buffer = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buffer.append(line)
            i += 1
            continue

        # Table row
        if '|' in line and line.strip().startswith('|'):
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            # Skip separator rows (---|---)
            if all(re.match(r'^[-: ]+$', c) for c in cells if c):
                i += 1
                continue
            table_buffer.append(cells)
            in_table = True
            i += 1
            continue
        else:
            if in_table and table_buffer:
                add_markdown_table(doc, table_buffer)
                table_buffer = []
                in_table = False

        stripped = line.strip()

        # Skip empty
        if not stripped:
            i += 1
            continue

        # H1
        if stripped.startswith('# '):
            if skip_first_h1:
                skip_first_h1 = False
                i += 1
                continue
            add_h1(doc, stripped[2:])
            i += 1
            continue

        # H2
        if stripped.startswith('## '):
            if skip_first_h2:
                skip_first_h2 = False
                i += 1
                continue
            add_h2(doc, stripped[3:])
            i += 1
            continue

        # H3
        if stripped.startswith('### '):
            add_h3(doc, stripped[4:])
            i += 1
            continue

        # H4+
        if stripped.startswith('#### '):
            add_h3(doc, stripped[5:])
            i += 1
            continue

        # Horizontal rule
        if stripped.startswith('---'):
            p = doc.add_paragraph()
            run = p.add_run("─" * 80)
            run.font.color.rgb = IOPEX_GRAY
            run.font.size = Pt(7)
            i += 1
            continue

        # Checklist item
        if stripped.startswith('- [ ]') or stripped.startswith('- [x]') or stripped.startswith('- [X]'):
            add_checklist_item(doc, stripped[5:])
            i += 1
            continue

        # Bullet
        if stripped.startswith('- ') or stripped.startswith('* '):
            add_bullet(doc, stripped[2:])
            i += 1
            continue

        # Numbered list
        if re.match(r'^\d+\. ', stripped):
            text = re.sub(r'^\d+\. ', '', stripped)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(2)
            parts = re.split(r'\*\*(.+?)\*\*', text)
            for k, part in enumerate(parts):
                if not part:
                    continue
                run = p.add_run(part)
                run.font.size = Pt(10)
                if k % 2 == 1:
                    run.font.bold = True
            i += 1
            continue

        # Blockquote
        if stripped.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(4)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '12')
            left.set(qn('w:space'), '6')
            left.set(qn('w:color'), '005BB5')
            pBdr.append(left)
            pPr.append(pBdr)
            run = p.add_run(stripped[2:])
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = IOPEX_GRAY
            i += 1
            continue

        # Regular paragraph
        add_body(doc, stripped)
        i += 1

    # Flush any remaining table or code
    if in_table and table_buffer:
        add_markdown_table(doc, table_buffer)
    if in_code and code_buffer:
        add_code_block(doc, '\n'.join(code_buffer))


# ── Footer ─────────────────────────────────────────────────────────────────
def add_footer(doc: Document, runbook_name: str):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"iOPEX Technologies  |  SHIFT Migration System  |  {runbook_name}  |  CONFIDENTIAL")
    run.font.size = Pt(7)
    run.font.color.rgb = IOPEX_GRAY


# ── Build All Runbooks ─────────────────────────────────────────────────────
RUNBOOKS = [
    {
        "md": "00_INDEX.md",
        "docx": "00_SHIFT_Migration_Index.docx",
        "meta": {
            "Document": "SHIFT Migration Runbook Index",
            "Classification": "CONFIDENTIAL",
            "Version": "1.0",
        }
    },
    {
        "md": "01_azure_deployment.md",
        "docx": "01_Azure_Deployment.docx",
        "meta": {
            "Scope": "Azure infrastructure deployment + container build",
            "Duration": "2–4 hours",
            "Prerequisite": "Azure subscription with Contributor access",
            "Classification": "CONFIDENTIAL",
        }
    },
    {
        "md": "02_stakeholder_dashboard.md",
        "docx": "02_Stakeholder_Dashboard.docx",
        "meta": {
            "Scope": "Activate live stakeholder dashboard — iOPEX + client visibility",
            "Duration": "30 minutes",
            "Prerequisite": "Azure deployment complete, deploy.sh run",
            "Audience": "iOPEX delivery team + Cisco migration stakeholders",
            "Classification": "CONFIDENTIAL",
        }
    },
    {
        "md": "03_openclaw_operations.md",
        "docx": "03_OpenClaw_Operations.docx",
        "meta": {
            "Scope": "OpenClaw shift-pmo agent — operating the migration via Option B (az CLI)",
            "Duration": "Reference document",
            "Prerequisite": "Azure deployment complete, TOOLS.md configured with APP_NAME",
            "Audience": "iOPEX delivery engineer (Jeremy Smith)",
            "Classification": "CONFIDENTIAL",
        }
    },
    {
        "md": "04_roles_responsibilities.md",
        "docx": "04_Roles_Responsibilities.docx",
        "meta": {
            "Scope": "Separation of duties — Lead Engineer (Accountable) vs Migration Engineer (Responsible)",
            "Applies to": "P0 through P7 — all migration phases",
            "Lead Engineer": "Jeremy Smith (iOPEX) — Accountable",
            "Migration Engineer": "[TBD] (iOPEX) — Responsible",
            "Classification": "CONFIDENTIAL",
        }
    },
    {
        "md": "P0_environment_setup.md",
        "docx": "P0_Environment_Setup.docx",
        "meta": {
            "Phase": "P0 — Environment Setup",
            "Duration": "1–2 days",
            "Prerequisite": "Azure deployment complete",
            "Classification": "CONFIDENTIAL",
        }
    },
    {
        "md": "P1_discovery.md",
        "docx": "P1_Discovery.docx",
        "meta": {
            "Phase": "P1 — Discovery, Gap Analysis & Permission Inventory",
            "Duration": "3–4 weeks",
            "Agents": "11 → 01 → 09 → 12 → 02 → 03",
            "Gate": "YC-P1 — Migration lead + security sign-off",
            "Classification": "CONFIDENTIAL",
        }
    },
    {
        "md": "P2_staging.md",
        "docx": "P2_Staging.docx",
        "meta": {
            "Phase": "P2 — Platform Validation & Staging",
            "Duration": "2–3 weeks",
            "Agents": "13 → 10",
            "Gate": "YC-P2 — 10/10 staging assertions",
            "Classification": "CONFIDENTIAL",
        }
    },
    {
        "md": "P3_safe_migration.md",
        "docx": "P3_Vault_Creation_Permissions.docx",
        "meta": {
            "Phase": "P3 — Vault Creation, Permissions & App Onboarding Setup",
            "Duration": "2–3 weeks",
            "Agents": "03 → 14",
            "Gate": "YC-P3 — Security + compliance sign-off",
            "Classification": "CONFIDENTIAL",
        }
    },
    {
        "md": "P4_pilot.md",
        "docx": "P4_Pilot_Migration.docx",
        "meta": {
            "Phase": "P4 — Pilot Migration (Wave 1)",
            "Duration": "1–2 weeks",
            "Agents": "04 → 05",
            "Gate": "YC-P4 — Heartbeat ≥ 95%, client sign-off",
            "Classification": "CONFIDENTIAL",
        }
    },
    {
        "md": "P5_production.md",
        "docx": "P5_Production_Migration.docx",
        "meta": {
            "Phase": "P5 — Production Migration (Waves 1–5)",
            "Duration": "4–6 weeks",
            "Agents": "04 → 05 → 06 → 14 → 07",
            "Gate": "YC-P5 — Per-wave gates + all waves complete",
            "Classification": "CONFIDENTIAL",
        }
    },
    {
        "md": "P6_parallel_running.md",
        "docx": "P6_Parallel_Running_Cutover.docx",
        "meta": {
            "Phase": "P6 — Parallel Running & Cutover",
            "Duration": "4–6 weeks",
            "Agents": "15 → 05 → 06 → 07",
            "Gate": "YC-P6 — 100% fleet migrated, CISO + client exec sign-off",
            "Classification": "CONFIDENTIAL",
        }
    },
    {
        "md": "P7_decommission.md",
        "docx": "P7_Decommission_Closeout.docx",
        "meta": {
            "Phase": "P7 — Decommission & Close-Out",
            "Duration": "2–3 weeks",
            "Agents": "07",
            "Gate": "YC-P7 — Final close-out sign-off",
            "Classification": "CONFIDENTIAL",
        }
    },
]


def build_runbook(rb: dict, script_dir: Path):
    md_path = script_dir / rb["md"]
    out_dir = script_dir / "docx"
    out_dir.mkdir(exist_ok=True)
    out_path = out_dir / rb["docx"]

    if not md_path.exists():
        print(f"  SKIP {rb['md']} — file not found")
        return

    print(f"  Building {rb['docx']}...")
    doc = create_document()
    add_footer(doc, rb["docx"].replace(".docx", "").replace("_", " "))
    parse_markdown_to_docx(doc, str(md_path), title_meta=rb.get("meta", {}))
    doc.save(str(out_path))
    print(f"  ✓ {out_path}")


def main():
    script_dir = Path(__file__).resolve().parent
    target = sys.argv[1].upper() if len(sys.argv) > 1 else None

    print("\niOPEX SHIFT Migration — Runbook Builder")
    print("=" * 50)

    built = 0
    for rb in RUNBOOKS:
        if target and target not in rb["md"].upper() and target not in rb["docx"].upper():
            continue
        build_runbook(rb, script_dir)
        built += 1

    print(f"\n{built} runbook(s) built → runbooks/docx/")
    print()


if __name__ == "__main__":
    main()
