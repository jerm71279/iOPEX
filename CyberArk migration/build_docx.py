#!/usr/bin/env python3
"""Build styled .docx documentation matching APM Terminals format.

Table style: dark navy headers (#1A365D) with white text, alternating row shading (#F7FAFC),
secondary tables with light blue headers (#E8F4F8). Arial font, 0.75in margins.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re

# ── Colors ───────────────────────────────────────────────────
NAVY = "1A365D"
DARK_BLUE = "2C5282"
MED_BLUE = "2B6CB0"
WHITE = "FFFFFF"
LIGHT_GRAY = "F7FAFC"
LIGHT_BLUE = "E8F4F8"
BODY_TEXT = "333333"
SUBTLE = "666666"


def keep_with_next(paragraph):
    """Set 'Keep with next' on a paragraph so it stays on the same page as the following element."""
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:keepNext {nsdecls("w")}/>'))


def keep_lines_together(paragraph):
    """Set 'Keep lines together' so the paragraph doesn't split across pages."""
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:keepLines {nsdecls("w")}/>'))


def set_row_cant_split(row):
    """Prevent a table row from splitting across a page break."""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    margins = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(margins)


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def add_styled_table(doc, headers, rows, header_bg=NAVY, header_fg=WHITE, alt_shading=True):
    """Add a table with APM Terminals styling.

    Applies cantSplit on every row to prevent rows from breaking across pages,
    and sets the header row to repeat at the top of each page if the table spans
    multiple pages.
    """
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    # Header row
    header_row = table.rows[0]
    set_row_cant_split(header_row)
    # Mark header row to repeat at top of each page
    trPr = header_row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

    for i, header_text in enumerate(headers):
        cell = header_row.cells[i]
        set_cell_shading(cell, header_bg)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        keep_lines_together(p)
        keep_with_next(p)
        run = p.add_run(header_text)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(10)
        if header_fg:
            run.font.color.rgb = RGBColor.from_string(header_fg)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        set_row_cant_split(row)
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_margins(cell)
            if alt_shading and r_idx % 2 == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            keep_lines_together(p)
            if r_idx == 0:
                keep_with_next(p)
            run = p.add_run(str(cell_text))
            run.font.name = "Arial"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(BODY_TEXT)

    return table


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    # Keep heading on the same page as the next element (table, paragraph, etc.)
    keep_with_next(h)
    keep_lines_together(h)
    if level == 1:
        # Each major section starts on a new page — prevents orphaned headings at bottom of page
        pPr = h._p.get_or_add_pPr()
        pPr.append(parse_xml(f'<w:pageBreakBefore {nsdecls("w")}/>'))
    for run in h.runs:
        run.font.name = "Arial"
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor.from_string(NAVY)
        elif level == 2:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
        elif level == 3:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor.from_string(MED_BLUE)
    return h


def add_para(doc, text, bold=False, size=10, color=BODY_TEXT, space_after=Pt(4), keep_next=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    p.paragraph_format.space_after = space_after
    if keep_next:
        keep_with_next(p)
        keep_lines_together(p)
    return p


def add_bullet(doc, text, bold=False, keep_next=True):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(BODY_TEXT)
    if keep_next:
        keep_with_next(p)
    return p


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    keep_lines_together(p)
    keep_with_next(p)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string("2D3748")
    # Light gray background via shading
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EDF2F7" w:val="clear"/>')
    p._p.get_or_add_pPr().append(shading)
    return p


def build():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    # ── Title Block ──────────────────────────────────────────
    p = doc.add_paragraph()
    run = p.add_run("CyberArk PAM Migration")
    run.font.name = "Arial"
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY)

    p = doc.add_paragraph()
    run = p.add_run("8-Agent AI Orchestrator System")
    run.font.name = "Arial"
    run.font.size = Pt(20)
    run.bold = True

    p = doc.add_paragraph()
    run = p.add_run("CyberArk On-Premises → Privilege Cloud")
    run.font.name = "Arial"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string(SUBTLE)

    p = doc.add_paragraph()
    run = p.add_run("Technical Documentation")
    run.font.name = "Arial"
    run.font.size = Pt(16)

    doc.add_paragraph()  # Spacer

    # ══════════════════════════════════════════════════════════
    # SECTION: Architecture
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "Architecture", 1)

    add_styled_table(doc,
        ["Component", "File", "Description"],
        [
            ["Orchestrator", "coordinator.py", "Sequences agents per phase, signal handlers, state management"],
            ["CLI", "cli.py", "User interface (status/start/run/advance/agent/preflight/agents)"],
            ["Agent Base", "core/base.py", "AgentBase ABC + AgentResult dataclass (status validation, approval timeout)"],
            ["State Machine", "core/state.py", "Atomic writes, fcntl file locking, backup recovery, raw data separation"],
            ["Audit Logger", "core/logging.py", "SHA-256 hash chain, SIEM-ready JSONL format"],
            ["On-Prem Client", "core/cyberark_client.py", "CyberArk PVWA REST client (context manager, env var credentials)"],
            ["Cloud Client", "core/cloud_client.py", "Privilege Cloud REST client (OAuth2 + legacy auth)"],
            [".gitignore", ".gitignore", "Excludes config.json, credentials, output/, __pycache__/"],
        ]
    )

    # ══════════════════════════════════════════════════════════
    # SECTION: Agent Registry
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "Agent Registry", 1)

    add_styled_table(doc,
        ["Agent", "Class", "Phases", "Key Capabilities"],
        [
            ["01-discovery", "DiscoveryAgent", "P1", "Applications API, multi-signal NHI detection, integration scanning"],
            ["02-gap-analysis", "GapAnalysisAgent", "P1", "10-domain assessment, 4-framework compliance mapping, scoring"],
            ["03-permissions", "PermissionMappingAgent", "P1, P3", "22 individual permissions via Safe Members API, sensitive flagging"],
            ["04-etl", "ETLOrchestrationAgent", "P4, P5", "Real API calls, watchdog timer, linked accounts, emergency unfreeze"],
            ["05-heartbeat", "HeartbeatAgent", "P4-P6", "10 post-migration checks (count, heartbeat, permissions, etc.)"],
            ["06-integration", "IntegrationRepointingAgent", "P5, P6", "CCP/AAM code scanning, replacement code generation"],
            ["07-compliance", "ComplianceAgent", "P5, P7", "PCI-DSS, NIST 800-53, HIPAA, SOX control mapping"],
            ["08-runbook", "RunbookAgent", "All", "Phase gate management, human approval workflow"],
        ]
    )

    # ══════════════════════════════════════════════════════════
    # SECTION: Phase Execution Map
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "Phase Execution Map", 1)

    add_styled_table(doc,
        ["Phase", "Name", "Agents", "Human Gate"],
        [
            ["P0", "Environment Setup", "(manual)", "—"],
            ["P1", "Discovery & Dependency Mapping", "01, 02, 03", "Review discovery + gaps + permissions"],
            ["P2", "Infrastructure Preparation", "(manual)", "—"],
            ["P3", "Safe & Policy Migration", "03", "Approve safe/policy plan"],
            ["P4", "Pilot Migration", "04, 05", "Approve pilot results"],
            ["P5", "Production Batches", "04, 05, 06, 07", "Approve each wave"],
            ["P6", "Parallel Running & Cutover", "05, 06, 07", "Approve cutover"],
            ["P7", "Decommission & Close-Out", "07", "Final sign-off"],
        ]
    )

    # ══════════════════════════════════════════════════════════
    # SECTION: Credential Management
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "Credential Management", 1)

    add_para(doc, "Credentials should be provided via environment variables (recommended) rather than stored in config.json.", keep_next=True)

    add_styled_table(doc,
        ["Environment Variable", "Purpose"],
        [
            ["CYBERARK_USERNAME", "On-prem PVWA service account username"],
            ["CYBERARK_PASSWORD", "On-prem PVWA service account password"],
            ["PCLOUD_CLIENT_ID", "Privilege Cloud OAuth2 client ID"],
            ["PCLOUD_CLIENT_SECRET", "Privilege Cloud OAuth2 client secret"],
            ["PCLOUD_IDENTITY_URL", "CyberArk Identity URL (override for OAuth2)"],
            ["PCLOUD_USERNAME", "Privilege Cloud username (legacy auth only)"],
            ["PCLOUD_PASSWORD", "Privilege Cloud password (legacy auth only)"],
        ]
    )

    add_para(doc, "Passwords are zeroed from memory immediately after the authentication attempt. On token expiry (HTTP 401), credentials are re-read from environment variables for re-authentication.", size=9)

    # ══════════════════════════════════════════════════════════
    # SECTION: Authentication
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "Authentication", 1)

    add_heading(doc, "On-Prem PVWA", 2)
    add_styled_table(doc,
        ["Auth Type", "Endpoint", "Notes"],
        [
            ["CyberArk", "/Auth/CyberArk/Logon", "Local vault authentication"],
            ["LDAP", "/Auth/LDAP/Logon", "Active Directory (recommended)"],
            ["RADIUS", "/Auth/RADIUS/Logon", "Multi-factor authentication"],
            ["Windows", "/Auth/Windows/Logon", "Windows integrated auth"],
            ["SAML", "—", "Not supported (requires browser redirect flow)"],
        ]
    )

    add_heading(doc, "Privilege Cloud (OAuth2 — Recommended)", 2)
    add_para(doc, "Modern Privilege Cloud tenants use OAuth2 via CyberArk Identity Security Platform:", keep_next=True)
    add_code_block(doc, "Token URL: {identity_url}/oauth2/platformtoken\nGrant type: client_credentials\nIdentity URL format: https://<tenant>.id.cyberark.cloud")

    add_heading(doc, "Privilege Cloud (Legacy)", 2)
    add_para(doc, "Set auth_method to 'legacy' in config. Uses /Auth/CyberArk/Logon on the Privilege Cloud PVWA URL.")

    # ══════════════════════════════════════════════════════════
    # SECTION: Configuration
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "Configuration", 1)

    add_heading(doc, "config.json", 2)
    add_styled_table(doc,
        ["Key", "Purpose", "Default"],
        [
            ["cyberark_on_prem.base_url", "On-prem PVWA URL", "—"],
            ["cyberark_on_prem.auth_type", "CyberArk / LDAP / RADIUS / Windows", "LDAP"],
            ["cyberark_on_prem.verify_ssl", "SSL certificate verification", "true"],
            ["cyberark_on_prem.timeout", "Request timeout (seconds)", "30"],
            ["cyberark_on_prem.rate_limit", "Seconds between requests", "0.1"],
            ["cyberark_on_prem.batch_size", "Pagination size (max 1000)", "1000"],
            ["privilege_cloud.base_url", "Privilege Cloud tenant URL", "—"],
            ["privilege_cloud.auth_method", "oauth2 (modern) or legacy", "oauth2"],
            ["privilege_cloud.identity_url", "CyberArk Identity URL for OAuth2", "—"],
            ["privilege_cloud.batch_size", "Import pagination size (max 1000)", "500"],
            ["servicenow.instance_url", "ServiceNow instance URL", "—"],
            ["output_dir", "Output directory", "./output"],
            ["log_level", "Logging level", "INFO"],
            ["environment", "Environment context", "dev"],
        ]
    )

    # ══════════════════════════════════════════════════════════
    # SECTION: Security Features
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "Security Features", 1)

    # Use light blue header style for this one
    add_styled_table(doc,
        ["Feature", "Implementation"],
        [
            ["Secrets Management", "Env var credentials, passwords zeroed from memory after auth, error sanitization strips secrets"],
            ["Atomic State Writes", "temp file → fsync → os.replace (POSIX atomic) + fcntl.flock(LOCK_EX) + .bak backup"],
            ["Signal Handlers", "SIGTERM/SIGINT save state and exit cleanly (exit code 128 + signal)"],
            ["Tamper-Evident Logging", "SHA-256 hash chain in every audit event — tampering breaks the chain"],
            ["Watchdog Timer", "threading.Timer auto-unfreezes all frozen accounts on ETL timeout (default 120 min)"],
            ["Approval Timeouts", "30-minute default, auto-deny in non-interactive mode (no TTY)"],
            ["HTTPS Enforcement", "Privilege Cloud connections reject non-HTTPS URLs"],
            ["Error Sanitization", "_safe_error() strips passwords, secrets, tokens from HTTP error messages"],
            [".gitignore", "Excludes config.json, *.pem, *.key, .env, output/, __pycache__/"],
        ],
        header_bg=LIGHT_BLUE,
        header_fg=None,
        alt_shading=False,
    )

    # ══════════════════════════════════════════════════════════
    # SECTION: Permission Mapping (Agent 03)
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "Permission Mapping (Agent 03)", 1)

    add_para(doc, "Privilege Cloud uses the same individual permission model as on-prem CyberArk. Agent 03 maps all 22 safe member permissions directly via the Safe Members API — there is no role-based simplification or permission loss.", keep_next=True)

    add_heading(doc, "All 22 CyberArk Safe Permissions (Mapped 1:1)", 2)
    add_code_block(doc,
        "UseAccounts              RetrieveAccounts         ListAccounts\n"
        "AddAccounts              UpdateAccountContent     UpdateAccountProperties\n"
        "InitiateCPMAccountManagementOperations            SpecifyNextAccountContent\n"
        "RenameAccounts           DeleteAccounts           UnlockAccounts\n"
        "ManageSafe               ManageSafeMembers        BackupSafe\n"
        "ViewAuditLog             ViewSafeMembers          AccessWithoutConfirmation\n"
        "CreateFolders            DeleteFolders            MoveAccountsAndFolders\n"
        "RequestsAuthorizationLevel1                       RequestsAuthorizationLevel2"
    )

    add_heading(doc, "Security-Sensitive Permissions (Flagged for Review)", 2)
    add_styled_table(doc,
        ["Permission", "Risk"],
        [
            ["ManageSafe", "Full safe admin — can modify safe properties"],
            ["ManageSafeMembers", "Can grant/revoke access to other users"],
            ["AccessWithoutConfirmation", "Bypasses dual-control approval workflow"],
            ["SpecifyNextAccountContent", "Can set the next password value"],
            ["InitiateCPMAccountManagementOperations", "Can trigger CPM verify/change/reconcile"],
            ["RequestsAuthorizationLevel1", "Dual-control authorization level 1"],
            ["RequestsAuthorizationLevel2", "Dual-control authorization level 2"],
        ]
    )

    add_para(doc, "Built-in members (Master, Batch, Backup Users, DR Users, Auditors, Operators) are automatically skipped. Phase P1 produces the analysis report. Phase P3 applies permissions via add_safe_member() / update_safe_member(), handling 409 conflicts by falling back to update.", size=9)

    # ══════════════════════════════════════════════════════════
    # SECTION: ETL Pipeline (Agent 04)
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "ETL Pipeline (Agent 04)", 1)

    add_para(doc, "Each batch executes 7 steps in sequence with real API calls:", keep_next=True)

    add_styled_table(doc,
        ["Step", "Action", "API Call"],
        [
            ["1. FREEZE", "Disable CPM automatic management", "PATCH /Accounts/{id} (automaticManagementEnabled=false)"],
            ["2. EXPORT", "Pull full account details + retrieve passwords", "GET /Accounts/{id} + POST .../Password/Retrieve"],
            ["3. TRANSFORM", "Map fields, preserve all properties", "Local transformation (no API)"],
            ["4. SAFE CREATION", "Create target safes if they don't exist", "POST /Safes (ManagingCPM, retention defaults)"],
            ["5. IMPORT", "Push accounts with retry + exponential backoff", "POST /Accounts (rate-limited)"],
            ["6. HEARTBEAT", "Trigger password verification", "POST /Accounts/{id}/Verify"],
            ["7. UNFREEZE", "Re-enable CPM management", "PATCH /Accounts/{id} (automaticManagementEnabled=true)"],
        ]
    )

    add_heading(doc, "Implementation Details", 2)
    add_bullet(doc, "Password retrieval: Accounts where retrieval fails are skipped (not imported with empty secrets)")
    add_bullet(doc, "Linked accounts: Logon/reconcile relationships recreated via POST /Accounts/{id}/LinkAccount")
    add_bullet(doc, "Field preservation: platformAccountProperties, remoteMachinesAccess, secretType, secretManagement")
    add_bullet(doc, "Failure threshold: >10% batch failure = 'failed', otherwise 'partial'")
    add_bullet(doc, "Both clients use context managers (with ... as client:) to prevent connection leaks")
    add_bullet(doc, "Emergency unfreeze: On any exception, all frozen accounts are automatically unfrozen", keep_next=False)

    add_heading(doc, "Wave Classification (5 Waves by Risk)", 2)
    add_styled_table(doc,
        ["Wave", "Risk", "Accounts"],
        [
            ["1", "LOW", "Test/Dev/Sandbox"],
            ["2", "MEDIUM", "Standard user accounts"],
            ["3", "MEDIUM-HIGH", "Infrastructure/network/admin"],
            ["4", "HIGH", "NHIs without CCP/AAM"],
            ["5", "CRITICAL", "NHIs with CCP/AAM (uses Applications API data)"],
        ]
    )

    # ══════════════════════════════════════════════════════════
    # SECTION: Validation Checks (Agent 05)
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "Validation Checks (Agent 05)", 1)

    add_styled_table(doc,
        ["#", "Check", "Description"],
        [
            ["1", "Count comparison", "Source vs target account counts within threshold"],
            ["2", "Heartbeat status", "All imported accounts have valid credentials"],
            ["3", "Permission mapping", "Translations applied correctly, exceptions flagged"],
            ["4", "Folder structure", "Safe-to-folder hierarchy preserved"],
            ["5", "Metadata integrity", "Descriptions and custom fields intact"],
            ["6", "Group assignments", "Group/role memberships translated"],
            ["7", "Password policies", "Rotation policies applied in target"],
            ["8", "Access patterns", "No unexpected permission escalations"],
            ["9", "Audit continuity", "Audit trail preserved across migration"],
            ["10", "Recording preservation", "PSM recordings archived"],
        ]
    )

    # ══════════════════════════════════════════════════════════
    # SECTION: Compliance Frameworks (Agent 07)
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "Compliance Frameworks (Agent 07)", 1)

    add_para(doc, "Maps migration actions to controls in 4 frameworks using data-driven check definitions (no lambdas — serializable config):", keep_next=True)

    add_styled_table(doc,
        ["Framework", "Control Groups"],
        [
            ["PCI-DSS v4.0", "Access control (8.x), Audit trail (10.x), Change management (6.x)"],
            ["NIST 800-53 Rev5", "IA-2/4/5, AC-2/3/6, AU-2/3/6/11, CA-7, SI-4"],
            ["HIPAA Security Rule", "164.312(a)(1), 164.312(b), 164.312(c)(1), 164.312(d)"],
            ["SOX IT Controls", "CC6.1-3, CC7.1-2, CC8.1"],
        ]
    )

    # ══════════════════════════════════════════════════════════
    # SECTION: NHI Detection
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "NHI Detection (Multi-Signal)", 1)

    add_styled_table(doc,
        ["Signal", "Priority", "Method"],
        [
            ["Platform-based", "Highest", "UnixSSHKeys, WinServiceAccount, WinScheduledTask, AzureServicePrincipal, AWSAccessKeys, HashiCorpVault"],
            ["Name patterns", "Medium", "^svc[_-], ^app[_-], ^api[_-], ^bot[_-], ^sys[_-], service.?account, daemon, scheduler"],
            ["Safe name patterns", "Lower", "appcred, servicecred, automation, cicd, pipeline, appidentit, machineidentit"],
        ],
        header_bg=LIGHT_BLUE,
        header_fg=None,
    )

    # ══════════════════════════════════════════════════════════
    # SECTION: API Endpoints
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "CyberArk REST API Endpoints", 1)

    add_heading(doc, "On-Prem PVWA (Source)", 2)
    add_styled_table(doc,
        ["Endpoint", "Method", "Purpose"],
        [
            ["/Auth/{type}/Logon", "POST", "Authentication (CyberArk/LDAP/RADIUS/Windows)"],
            ["/Auth/Logoff", "POST", "Session termination"],
            ["/Accounts", "GET", "Account enumeration (paginated, count-based)"],
            ["/Accounts/{id}", "GET", "Full account details (includes linked accounts)"],
            ["/Accounts/{id}", "PATCH", "Update account properties (JSON Patch)"],
            ["/Accounts/{id}/Password/Retrieve", "POST", "Password retrieval"],
            ["/Safes", "GET", "Safe enumeration (system safes filtered)"],
            ["/Safes/{name}/Members", "GET", "Safe member/permission enumeration"],
            ["/Platforms/Targets", "GET", "Platform enumeration (v12+)"],
            ["/Platforms/{id}/Export", "POST", "Platform export (ZIP package)"],
            ["/Applications", "GET", "CCP/AAM application enumeration"],
            ["/Applications/{id}/Authentications", "GET", "Application auth methods"],
            ["/Activities", "GET", "Audit log retrieval"],
            ["/Server/Verify", "GET", "Version and health info"],
        ]
    )

    add_heading(doc, "Privilege Cloud (Target)", 2)
    add_styled_table(doc,
        ["Endpoint", "Method", "Purpose"],
        [
            ["{identity_url}/oauth2/platformtoken", "POST", "OAuth2 token (client_credentials)"],
            ["/Auth/CyberArk/Logon", "POST", "Legacy authentication"],
            ["/Safes", "POST", "Safe creation"],
            ["/Safes/{name}/Members", "POST", "Add safe member with individual permissions"],
            ["/Safes/{name}/Members/{member}", "PUT", "Update safe member permissions"],
            ["/Accounts", "POST", "Account import"],
            ["/Accounts/{id}/Verify", "POST", "Password verification (heartbeat)"],
            ["/Accounts/{id}/Password/Retrieve", "POST", "Password retrieval"],
            ["/Accounts/{id}/LinkAccount", "POST", "Link logon/reconcile account"],
            ["/Accounts/{id}", "DELETE", "Account deletion (rollback)"],
            ["/Platforms/Import", "POST", "Platform package import (ZIP)"],
        ]
    )

    # ══════════════════════════════════════════════════════════
    # SECTION: State Machine
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "State Machine", 1)

    add_para(doc, "Migration state persists to output/state/migration_state.json via atomic writes:", keep_next=True)

    add_styled_table(doc,
        ["Feature", "Implementation"],
        [
            ["Persistence", "JSON file with atomic writes (temp → fsync → os.replace)"],
            ["Concurrency", "fcntl.flock(LOCK_EX) prevents concurrent writes"],
            ["Backup", ".bak file created before every write; recovery on corruption"],
            ["Raw Data", "Large datasets stored in output/state/raw/ (not in main state file)"],
            ["List Caps", "Steps: 5000, Errors: 1000, Approvals: 500"],
            ["Resume", "Coordinator picks up from last completed step after restart"],
        ],
        header_bg=LIGHT_BLUE,
        header_fg=None,
        alt_shading=False,
    )

    # ══════════════════════════════════════════════════════════
    # SECTION: Test Environment Decision Point
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "Test Environment & Sandbox (Next Steps Decision Point)", 1)

    add_para(doc, "The system currently has no dedicated virtual test environment. Before running the ETL pipeline for the first time, a decision is needed on the testing approach.", keep_next=True)

    add_heading(doc, "Current Capabilities", 2)
    add_styled_table(doc,
        ["Capability", "What It Does", "Limitations"],
        [
            ["--dry-run mode", "Skips agent execution, prints what would run", "No data flow simulation"],
            ["environment config", "Stamps dev/staging/prod into audit logs", "Metadata only — no behavior change"],
            ["P4 Pilot (Wave 1)", "Limits first batch to 50 test/dev accounts", "Hits real APIs — not a sandbox"],
        ],
        header_bg=LIGHT_BLUE,
        header_fg=None,
    )

    add_heading(doc, "Decision Matrix", 2)
    add_styled_table(doc,
        ["Option", "Description", "Effort", "Fidelity"],
        [
            ["A. Lab Instance", "CyberArk PVWA dev/lab + Privilege Cloud test tenant", "Low", "High"],
            ["B. Mock API Server", "Flask/FastAPI emulating PVWA/Cloud APIs", "Medium", "Medium"],
            ["C. Unit Test Fixtures", "pytest with mocked API clients", "Medium", "Medium"],
            ["D. Docker Harness", "Mock API containers with pre-seeded data", "High", "High"],
            ["E. Hybrid (Recommended)", "A for integration + C for CI/CD regression", "Medium", "High"],
        ]
    )

    add_heading(doc, "Recommended: Hybrid (Option E)", 2)
    add_bullet(doc, "Integration testing: Separate config files per environment (config.dev.json, config.prod.json)")
    add_bullet(doc, "CI/CD regression: pytest fixtures with mocked API responses")
    add_bullet(doc, "Wave 1 is the built-in canary — test/dev/sandbox accounts migrated first as a safe first-run", keep_next=False)

    # ══════════════════════════════════════════════════════════
    # SECTION: CLI Commands
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "CLI Commands", 1)

    add_styled_table(doc,
        ["Command", "Description"],
        [
            ["python3 cli.py status", "Show migration status and phase breakdown"],
            ["python3 cli.py start <id>", "Start new migration with given ID"],
            ["python3 cli.py run <P#>", "Run all agents for a phase"],
            ["python3 cli.py run <P#> --dry-run", "Simulate phase without API calls"],
            ["python3 cli.py advance", "Advance to next phase"],
            ["python3 cli.py agent <key>", "Run a single agent (e.g., 01-discovery)"],
            ["python3 cli.py agent <key> --phase P#", "Run agent for a specific phase"],
            ["python3 cli.py preflight", "Run all agent preflight checks"],
            ["python3 cli.py agents", "List all available agents"],
        ]
    )

    # ══════════════════════════════════════════════════════════
    # SECTION: Audit Logging
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "Audit Logging", 1)

    add_para(doc, "All agent actions are logged as structured JSON events (JSONL) to output/logs/<agent_id>.audit.jsonl. Every event includes a chain_hash — a running SHA-256 hash linking each event to its predecessor. Tampering with log entries breaks the chain, making modifications detectable during compliance audits.", keep_next=True)

    add_code_block(doc,
        '{\n'
        '  "timestamp": "2026-03-03T03:16:47.853Z",\n'
        '  "agent_id": "agent_01_discovery",\n'
        '  "environment": "prod",\n'
        '  "action": "discovery_start",\n'
        '  "details": {"phase": "P1"},\n'
        '  "user": "maverick",\n'
        '  "session_id": "6a410ecc",\n'
        '  "result": "success",\n'
        '  "chain_hash": "a3f8c2...e91d04"\n'
        '}'
    )

    # ── Save ─────────────────────────────────────────────────
    out_path = "/home/maverick/projects/iOPEX/CyberArk migration/CyberArk_Migration_Agent_System_Documentation.docx"
    doc.save(out_path)
    print(f"Saved to {out_path}")

    # Also copy to OneDrive
    import shutil
    onedrive = "/mnt/c/Users/1/OneDrive/Documents/iOPEX/CyberArk_Migration_Agent_System_Documentation.docx"
    shutil.copy2(out_path, onedrive)
    print(f"Copied to {onedrive}")


if __name__ == "__main__":
    build()
